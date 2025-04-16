# auto_report.py

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import sqlite3

from PIL import Image
import subprocess



# Configuration
project_dir = os.getcwd()
lab_number_path = os.path.join(project_dir, 'report/lab-number.txt')
scr_path = os.path.join(project_dir, 'src')
test_dir = os.path.join(project_dir, 'tests')
diagram_path = os.path.join(project_dir, 'report/diagram')
git_link = 'your-link'

# Change path to your own
database_path = '!project-root-path!/database/labs.db'

packy_dir = '!project-root-path!/../packy/'

# Create a new Document
doc = Document()

#lab_number
#condition_img_path
#condition_text

# Get lab data from the database by lab_number
def get_lab_data(lab_number):
    try:
        # Ensure lab_number is a float for matching with the database
          # Strip any whitespace and convert to float

        # Connect to the database
        conn = sqlite3.connect(database_path)  # Use your actual database path
        cursor = conn.cursor()

        # Query to select data for the given lab_number
        cursor.execute('''
            SELECT title, purpose, tasks, task_image_path FROM report_data WHERE lab_number = ?
        ''', (lab_number,))

        # Fetch the data
        data = cursor.fetchone()

        # Debugging: print out the result
        if data:
            print(f"Lab data found for lab_number {lab_number}: {data}")
        else:
            print(f"No data found for lab_number: {lab_number}")

        # Close the connection
        conn.close()

        # Return the data if found, otherwise return None
        if data:
            return {
                'lab_number': lab_number,
                'title': data[0],
                'purpose': data[1],
                'tasks': data[2],
                'task_image_path': data[3]
            }
        else:
            return None

    except sqlite3.Error as e:
        print(f"SQLite error: {e}")
        return None
    except ValueError as e:
        print(f"Error converting lab_number to float: {e}")
        return None



# Get number of lab
with open(lab_number_path, 'r', encoding='utf-8') as file:
    lab_number = file.read()
    print(lab_number)

# Get data for report
lab_data           = get_lab_data(lab_number)


# Save data categorized only if lab_data is not None
if lab_data is not None:
    lab_theme_name     = lab_data['title']
    lab_goal           = lab_data['purpose']
    condition_text     = lab_data['tasks']
    condition_img_path = packy_dir + lab_data['task_image_path']
    
    print(condition_img_path)

    # Proceed with creating the document and report texts
    name_of_subject = "Об’єктно-орієнтоване програмування"
    name_of_group = "Група"
    my_name = ["Прізвище", "Ім'я", "Побатькове"]
    doc_name = f'ООП_{name_of_group}_{my_name[0]}_ЛР-{lab_number}.docx' 

    conclusion_text = f'У результаті виконання лабораторної роботи я зміг {lab_goal}'
    report_header_text = (
            'ЗВІТ\n'
            f'про виконання лабораторної роботи № {lab_number}\n'
            f'{lab_theme_name}\n'
            'з дисципліни\n'
            f'"{name_of_subject}"\n'
            f'студента групи {name_of_group}\n'
            f'{my_name[0]} {my_name[1]} {my_name[2]}\n'
            )
else:
    print("Error: lab_data is None.")
    # Optionally exit or raise an exception here if you can't proceed without the data.
    exit(1)  # Stop execution if there's no data for the lab.



def is_image(file_path):
    try:
        with Image.open(file_path) as img:
            img.verify()  # Check if it's a valid image
        return True
    except Exception:
        return False



# Start of the doc by header
def doc_header():

    print("Start creating report file...")

    ######################
    ## Header of report ##
    ######################

    print("Adding header of report...")

        
    report_header           = doc.add_paragraph()
    report_header_run       = report_header.add_run(report_header_text)
    report_header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    report_header_run.font.name = 'Times New Roman Cyr'
    report_header_run.font.size = Pt(14)
    report_header_run.bold      = True

    doc.add_page_break()


# Adding task condition to the doc
def doc_task():

    ####################
    ## Task Condition ##
    #################### 

    print("Adding task condition...")

    doc.add_heading(f'Мета: {lab_goal}')
    doc.add_heading('Умова завдання:', level=1)


    task_cond_pgh           = doc.add_paragraph()

    task_cond_run           = task_cond_pgh.add_run(condition_text)
    task_cond_run.font.name = 'Times New Roman Cyr'
    task_cond_run.font.size = Pt(12)

    
    doc.add_picture(condition_img_path, width=Inches(5))

    doc.add_page_break()


# Adding structure && UML diagram imgs
def doc_diagram():
    
    ############################## 
    ## Structure && UML Diagram ##
    ############################## 

    print("Adding structure diagram...")

    doc.add_heading('Діаграми програми:', level=1)

    # add all diagrams in directory
    for imgname in os.listdir(diagram_path):
        # skip if it's not a file
        img_path = os.path.join(diagram_path, imgname)
        if is_image(img_path):
            print(f"Adding {imgname}...")  # Output name of file
            # add image
            doc.add_picture(img_path, width=Inches(5))
            doc.add_page_break()


# Function for adding programme files from directory
def doc_script(folder_path):
    for filename in os.listdir(folder_path):
        # skip if it's not a file
        file_path = os.path.join(folder_path, filename)
        if os.path.isfile(file_path):
            print(f"Adding {filename}...")  # Output name of file
            doc.add_heading(filename, level=2)

            # Reading file
            with open(file_path, 'r', encoding='utf-8') as file:
                file_text = file.read()

            # Adding text of the file to the document
            file_paragraph     = doc.add_paragraph()

            file_run           = file_paragraph.add_run(file_text)
            file_run.font.name = 'Courier New'
            file_run.font.size = Pt(10)

            
            doc.add_page_break()


# Adding github repo link to doc
def doc_git():
    
    ##############
    ## Git Repo ##
    ############## 

    print("Adding Git repo link...")

    doc.add_heading('Посилання на git-репозиторій з проктом', level=2)


    git_link_pgh           = doc.add_paragraph() 

    git_link_run           = git_link_pgh.add_run(git_link)
    git_link_run.font.name = 'Courier New'
    git_link_run.font.size = Pt(10)


    doc.add_page_break()


# Adding unit-test and it results
def doc_tests():
    
    ################ 
    ## Unit Tests ##
    ################

    print("Adding unit tests...")

    doc.add_heading('Unit-тест:')
    doc_script(test_dir)

    doc.add_heading('Результати unit-тесту:')

    # Run tests and add result to the doc
    
    try:
        result = subprocess.run(
            ["ctest", "--test-dir", "build", "--output-on-failure", "-j12"],
            capture_output=True,
            text=True,
            check=True
        )

        tests_result = doc.add_paragraph()

        # Extract test output from result
        test_output = result.stdout if result.stdout else "No output from tests."

        tests_results_run = tests_result.add_run(test_output)
        tests_results_run.font.name = 'Courier New'
        tests_results_run.font.size = Pt(10)

    except subprocess.CalledProcessError as e:
        print(f"Error running tests: {e}")

    doc.add_page_break()


# Adding conclusion to the doc
def doc_conclusion():

    #################
    ## Conclusions ##
    #################

    print("Adding conclusion...")

    doc.add_heading('Висновки:')


    conclusion_pgh           = doc.add_paragraph()

    conclusion_run           = conclusion_pgh.add_run(conclusion_text)
    conclusion_run.font.name = 'Times New Roman Cyr'
    conclusion_run.font.size = Pt(12)

# Generate report with this structure
def gen_doc():
    doc_header()
    doc_task()
    doc_diagram()
    doc_script(scr_path)
    doc_git()
    doc_tests()
    doc_conclusion()

    # Save the Document
    doc.save(f'{project_dir}/{doc_name}')
    print(f"\nCreated {doc_name}")


